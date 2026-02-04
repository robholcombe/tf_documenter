# terraform_documenter/docx_generator.py

from PIL import Image
import os
from docx import Document
from docx.shared import Inches
from typing import Dict, Any
from .utils import PROVIDER_DISPLAY_NAMES

class DocxGenerator:
    def __init__(self, data: Dict[str, Any], provider: str, directory: str, conceptual_img: str, networking_img: str):
        self.data = data
        self.provider = provider
        self.provider_name = PROVIDER_DISPLAY_NAMES.get(provider, provider.upper())
        self.directory = directory
        self.conceptual_img = conceptual_img
        self.networking_img = networking_img
        self.document = Document()
        # Generate a summary of key components to be used dynamically
        self.summary_points = self._generate_summary_points()

    def create_sdd(self, output_filename: str):
        """Main method to generate the complete SDD document."""
        self.document.add_heading(f'Solution Design Document: {self.provider_name} Architecture', level=0)
        
        self._add_introduction()
        self._add_executive_summary()
        self._add_conceptual_architecture()
        self._add_networking_design()
        self._add_compute_and_storage()
        self._add_security_and_identity()
        self._add_terraform_environment()
        self._add_appendix()

        self.document.save(output_filename)
        print(f"Successfully generated SDD: {output_filename}")

    def _generate_summary_points(self) -> Dict[str, int]:
        """Helper method to count key resources for dynamic summaries."""
        points = {}
        resources = self.data.get("resources", {})
        
        # Consolidate resource counts across different provider naming conventions
        points["vpcs"] = len(resources.get("vpcs", [])) + len(resources.get("vnets", [])) + len(resources.get("networks", []))
        points["instances"] = len(resources.get("instances", [])) + len(resources.get("vms", []))
        points["databases"] = len(resources.get("rds", [])) + len(resources.get("sql_db", [])) + len(resources.get("sql", [])) + len(resources.get("dynamodb", [])) + len(resources.get("cosmosdb", []))
        points["containers"] = len(resources.get("eks", [])) + len(resources.get("ecs", [])) + len(resources.get("aks", [])) + len(resources.get("gke", []))
        points["functions"] = len(resources.get("lambda", [])) + len(resources.get("functions", []))
        points["storage"] = len(resources.get("s3", [])) + len(resources.get("storage_accounts", [])) + len(resources.get("storage_buckets", []))
        points["load_balancers"] = len(resources.get("alb", [])) + len(resources.get("nlb", [])) + len(resources.get("app_gateway", [])) + len(resources.get("load_balancer", [])) + len(resources.get("load_balancing", []))
        
        # Return only the categories with detected resources
        return {k: v for k, v in points.items() if v > 0}

    def _add_introduction(self):
        self.document.add_heading('1. Introduction', level=1)
        self.document.add_heading('1.1 Purpose of the Document', level=2)
        
        purpose_text = (
            f"This document describes the cloud architecture for the {self.provider_name} environment, as defined by the "
            f"Terraform configuration files. It provides a detailed overview of the provisioned infrastructure, "
            f"including networking, compute, storage, and security components based on the parsed resources."
        )
        self.document.add_paragraph(purpose_text)
        
        self.document.add_heading('1.2 Scope', level=2)
        
        # Dynamically calculate the scope based on parsed data
        resource_categories_found = len([cat for cat, res_list in self.data.get("resources", {}).items() if res_list])
        total_modules = len(self.data.get("modules", []))
        
        scope_text = (
            f"The scope of this document is limited to the resources defined in the Terraform files. "
            f"The analysis has identified {resource_categories_found} distinct resource categories "
            f"and {total_modules} module(s) used in the configuration."
        )
        self.document.add_paragraph(scope_text)

        self.document.add_heading('1.3 Detected Cloud Provider', level=2)
        self.document.add_paragraph(self.provider_name)
        self.document.add_heading('1.4 Source Terraform Directory', level=2)
        self.document.add_paragraph(self.directory)

    def _add_executive_summary(self):
        self.document.add_heading('2. Executive Summary', level=1)
        
        # Generate dynamic summary phrases based on found resources
        summary_phrases = []
        if self.summary_points.get("vpcs"):
            summary_phrases.append(f"a virtual network foundation with {self.summary_points['vpcs']} VPC/VNet(s)")
        if self.summary_points.get("instances"):
            summary_phrases.append(f"{self.summary_points['instances']} compute instance(s)")
        if self.summary_points.get("containers"):
            summary_phrases.append(f"containerized workloads on {self.summary_points['containers']} cluster(s)")
        if self.summary_points.get("functions"):
            summary_phrases.append(f"{self.summary_points['functions']} serverless function(s)")
        if self.summary_points.get("databases"):
            summary_phrases.append(f"{self.summary_points['databases']} managed database instance(s)")
        if self.summary_points.get("storage"):
            summary_phrases.append(f"object storage using {self.summary_points['storage']} bucket/account(s)")
        if self.summary_points.get("load_balancers"):
            summary_phrases.append(f"load balancing across {self.summary_points['load_balancers']} balancer(s)")

        if not summary_phrases:
            summary_text = (
                f"This document outlines the infrastructure for the {self.provider_name} environment. "
                f"No primary architectural components were detected in the Terraform configuration, which may indicate a focus "
                f"on other resource types such as IAM, security policies, or organizational roles."
            )
        else:
            # Build the summary sentence with proper grammar
            if len(summary_phrases) > 2:
                summary_sentence = ", ".join(summary_phrases[:-1]) + ", and " + summary_phrases[-1]
            else:
                summary_sentence = " and ".join(summary_phrases)
            
            summary_text = (
                f"This document outlines the infrastructure for the {self.provider_name} environment as defined by Terraform. "
                f"The architecture's key components include {summary_sentence}."
            )
        self.document.add_paragraph(summary_text)

    def _add_conceptual_architecture(self):
        self.document.add_heading('3. Conceptual Architecture', level=1)
        self.document.add_heading('3.1 Architecture Overview', level=2)
        self.document.add_paragraph('The following diagram provides a high-level overview of the solution components and their interactions.')
        self.document.add_heading('3.2 Conceptual Architecture Diagram', level=2)
        try:
            img_path = self.conceptual_img
            img = Image.open(img_path)
            original_width, original_height = img.size
            aspect_ratio = original_width / original_height if original_height > 0 else 1

            max_height_inches = Inches(10 / 2.54)
            max_width_inches = Inches(16 / 2.54)

            scaled_height = max_height_inches
            scaled_width = scaled_height * aspect_ratio

            if scaled_width > max_width_inches:
                scaled_width = max_width_inches
                scaled_height = scaled_width / aspect_ratio if aspect_ratio > 0 else 0

            self.document.add_picture(img_path, height=scaled_height)
        except FileNotFoundError:
            self.document.add_paragraph('[Diagram file not found]')
        except Exception as e:
            self.document.add_paragraph(f'[Error adding diagram: {e}]')

    def _add_networking_design(self):
        self.document.add_heading('4. Networking Design', level=1)
        # ... (code for VPCs/VNets, Subnets, etc.) ...
        
        self.document.add_heading('4.4 Networking and Security Diagram', level=2)
        try:
            img_path = self.networking_img
            img = Image.open(img_path)
            original_width, original_height = img.size
            aspect_ratio = original_width / original_height if original_height > 0 else 1

            max_height_inches = Inches(10 / 2.54)
            max_width_inches = Inches(16 / 2.54)

            scaled_height = max_height_inches
            scaled_width = scaled_height * aspect_ratio

            if scaled_width > max_width_inches:
                scaled_width = max_width_inches
                scaled_height = scaled_width / aspect_ratio if aspect_ratio > 0 else 0

            self.document.add_picture(img_path, height=scaled_height)
        except FileNotFoundError:
            self.document.add_paragraph('[Diagram file not found]')
        except Exception as e:
            self.document.add_paragraph(f'[Error adding diagram: {e}]')
    
    def _add_compute_and_storage(self):
        self.document.add_heading('5. Compute and Storage', level=1)
        
        # --- Compute Instances ---
        instances = self.data['resources'].get('instances', [])
        if instances:
            self.document.add_heading('5.1 Compute Instances', level=2)
            table = self.document.add_table(rows=1, cols=3, style='Table Grid')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Name'
            hdr_cells[1].text = 'Machine Type'
            hdr_cells[2].text = 'Zone'
            for instance in instances:
                row_cells = table.add_row().cells
                row_cells[0].text = instance['name']
                row_cells[1].text = str(instance['config'].get('machine_type', ['N/A'])[0])
                row_cells[2].text = str(instance['config'].get('zone', ['N/A'])[0])
        
        # --- Storage Buckets ---
        buckets = self.data['resources'].get('storage_buckets', [])
        if buckets:
            self.document.add_heading('5.2 Storage Buckets', level=2)
            table = self.document.add_table(rows=1, cols=2, style='Table Grid')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Name'
            hdr_cells[1].text = 'Location'
            for bucket in buckets:
                row_cells = table.add_row().cells
                row_cells[0].text = bucket['name']
                row_cells[1].text = str(bucket['config'].get('location', ['N/A'])[0])
    
    def _add_security_and_identity(self):
        self.document.add_heading('6. Security and Identity', level=1)
        
        # --- Firewall Rules ---
        firewalls = self.data['resources'].get('firewall', [])
        if firewalls:
            self.document.add_heading('6.1 Firewall Rules', level=2)
            table = self.document.add_table(rows=1, cols=3, style='Table Grid')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Name'
            hdr_cells[1].text = 'Network'
            hdr_cells[2].text = 'Allowed Protocols/Ports'
            for rule in firewalls:
                row_cells = table.add_row().cells
                row_cells[0].text = rule['name']
                row_cells[1].text = str(rule['config'].get('network', ['N/A'])[0]).split('/')[-1] # Clean up network name
                
                # Summarize allow rules
                allow_rules = rule['config'].get('allow', [])
                rules_summary = []
                for allow in allow_rules:
                    protocol = allow.get('protocol')
                    ports = allow.get('ports', [])
                    rules_summary.append(f"{protocol}: {', '.join(ports)}")
                row_cells[2].text = '; '.join(rules_summary) if rules_summary else 'All'

        # --- Service Accounts ---
        service_accounts = self.data['resources'].get('iam_roles', [])
        if service_accounts:
            self.document.add_heading('6.2 Service Accounts', level=2)
            table = self.document.add_table(rows=1, cols=2, style='Table Grid')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Account ID'
            hdr_cells[1].text = 'Display Name'
            for sa in service_accounts:
                row_cells = table.add_row().cells
                row_cells[0].text = str(sa['config'].get('account_id', ['N/A'])[0])
                row_cells[1].text = str(sa['config'].get('display_name', ['N/A'])[0])

    def _add_terraform_environment(self):
        self.document.add_heading('7. Terraform Environment', level=1)
        self.document.add_heading('7.1 Terraform Variables', level=2)
        variables = self.data.get('variables', [])
        if variables:
            table = self.document.add_table(rows=1, cols=3, style='Table Grid')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Name'
            hdr_cells[1].text = 'Description'
            hdr_cells[2].text = 'Default Value'
            for var in variables:
                row_cells = table.add_row().cells
                row_cells[0].text = var['name']
                row_cells[1].text = str(var['description'])
                row_cells[2].text = str(var['default'])
    
    def _add_appendix(self):
        self.document.add_heading('Appendix', level=1)
        self.document.add_heading('A. Raw Resource List', level=2)
        for category, resources in self.data['resources'].items():
            if resources:
                self.document.add_paragraph(f"{category.replace('_', ' ').title()}:", style='List Bullet')
                for resource in resources:
                    self.document.add_paragraph(f"  - {resource['type']}.{resource['name']}", style='List 2')